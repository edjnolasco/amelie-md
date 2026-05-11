from pathlib import Path

from docx import Document

from amelie_md.exporters.docx import DocxExporter, DocxMetadata
from amelie_md.document import AmelieDocument  # ← IMPORT NUEVO


def test_docx_exporter_creates_docx_with_core_elements(tmp_path: Path) -> None:
    output_path = tmp_path / "output.docx"

    exporter = DocxExporter(
        metadata=DocxMetadata(
            title="Test Document",
            author="Amelie Suite",
            date="2026-04-28",
        )
    )

    # 🔁 NUEVO: documento estructurado (sin Markdown)
    document_model = AmelieDocument(
        blocks=[
            {"type": "heading", "level": 1, "text": "Main Title"},
            {"type": "heading", "level": 2, "text": "Section"},
            {"type": "paragraph", "text": "This is a **bold** paragraph with *italic*, `code` and [OpenAI](https://openai.com)."},
            {"type": "list_item", "text": "First item", "level": 0, "ordered": False},
            {"type": "list_item", "text": "Second item", "level": 0, "ordered": False},
            {
                "type": "table",
                "rows": [
                    ["Name", "Value"],
                    ["A", "1"],
                ],
            },
            {"type": "code", "code": 'print("hello")'},
        ]
    )

    # 🔁 NUEVO: export directo desde AmelieDocument
    exporter.export_document(document_model, output_path)

    assert output_path.exists()

    document = Document(output_path)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Test Document" in paragraphs
    assert "Main Title" in paragraphs
    assert "Section" in paragraphs
    assert "This is a bold paragraph with italic, code and OpenAI." in paragraphs
    assert "First item" in paragraphs
    assert "Second item" in paragraphs
    assert 'print("hello")' in paragraphs

    inline_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "This is a bold paragraph with italic, code and OpenAI." in paragraph.text
    )

    assert any(run.text == "bold" and run.bold for run in inline_paragraph.runs)
    assert any(run.text == "italic" and run.italic for run in inline_paragraph.runs)
    assert any(run.text == "code" for run in inline_paragraph.runs)
    document_xml = inline_paragraph._p.xml
    assert "w:hyperlink" in document_xml
    assert "OpenAI" in document_xml


    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Name"
    assert document.tables[0].cell(1, 1).text == "1"

def test_docx_exporter_renders_admonition_block(tmp_path: Path) -> None:
    output_path = tmp_path / "admonition.docx"

    exporter = DocxExporter()

    document_model = AmelieDocument(
        blocks=[
            {
                "type": "admonition",
                "kind": "warning",
                "title": "Warning",
                "text": "Be careful with **production** data.",
            }
        ]
    )

    exporter.export_document(document_model, output_path)

    document = Document(output_path)

    assert len(document.tables) >= 1

    cell_text = document.tables[0].cell(0, 0).text

    assert "Warning" in cell_text
    assert "Be careful with production data." in cell_text


def test_docx_exporter_renders_definition_block(tmp_path: Path) -> None:
    output_path = tmp_path / "definition.docx"

    exporter = DocxExporter()

    document_model = AmelieDocument(
        blocks=[
            {
                "type": "definition",
                "title": "Concepto clave",
                "text": "A **definition** explains a concept precisely.",
            }
        ]
    )

    exporter.export_document(document_model, output_path)

    document = Document(output_path)

    assert len(document.tables) >= 1

    cell_text = document.tables[0].cell(0, 0).text

    assert "Concepto clave" in cell_text
    assert "A definition explains a concept precisely." in cell_text


def test_docx_exporter_renders_quote_block(tmp_path: Path) -> None:
    output_path = tmp_path / "quote.docx"

    exporter = DocxExporter()

    document_model = AmelieDocument(
        blocks=[
            {
                "type": "quote",
                "title": "Autor",
                "text": "Knowledge is structured meaning.",
            }
        ]
    )

    exporter.export_document(document_model, output_path)

    document = Document(output_path)
    paragraphs = "\\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Knowledge is structured meaning." in paragraphs
    assert "— Autor" in paragraphs
