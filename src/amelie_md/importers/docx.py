from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from amelie_md.document import AmelieDocument, DocumentBlock


class DocxImporter:
    """Imports a .docx file into Amelie's intermediate document model."""

    HEADING_PATTERNS: tuple[tuple[re.Pattern[str], int], ...] = (
        (re.compile(r"^\s*cap[ií]tulo\s+\d+[\.:]?\s+.+$", re.IGNORECASE), 1),
        (re.compile(r"^\s*\d+\.\d+\.\d+\.?\s+\S.+$"), 3),
        (re.compile(r"^\s*\d+\.\d+\.?\s+\S.+$"), 2),
        (re.compile(r"^\s*\d+\.\s+\S.+$"), 1),
        (re.compile(r"^\s*secci[oó]n\s+\d+[\.:]?\s+.+$", re.IGNORECASE), 2),
        (re.compile(r"^\s*anexo\s+([a-z]|\d+)[\.:]?\s*.*$", re.IGNORECASE), 1),
    )

    MAX_HEADING_LENGTH = 120

    def import_file(self, input_path: str | Path) -> AmelieDocument:
        path = Path(input_path)

        if not path.exists():
            raise FileNotFoundError(f"Input DOCX file not found: {path}")

        if path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file, got: {path.suffix}")

        document = Document(path)
        blocks: list[DocumentBlock] = []

        for element in document.element.body:
            tag = element.tag.lower()

            if tag.endswith("}p"):
                paragraph = self._paragraph_from_element(document, element)
                if paragraph is not None:
                    blocks.append(paragraph)

            elif tag.endswith("}tbl"):
                table = self._table_from_element(document, element)
                if table is not None:
                    blocks.append(table)

        return AmelieDocument(blocks=blocks)

    def _paragraph_from_element(
        self,
        document: Document,
        element: Any,
    ) -> DocumentBlock | None:
        paragraph = self._lookup_paragraph(document, element)

        if paragraph is None:
            return None

        text = paragraph.text.strip()
        if not text:
            return None

        style_name = paragraph.style.name if paragraph.style else ""

        # 1. Heading real de Word
        style_heading_level = self._detect_heading_level_from_style(style_name)
        if style_heading_level is not None:
            return DocumentBlock(
                type="heading",
                text=text,
                level=style_heading_level,
            )

        # 2. Lista nativa Word (numPr)
        word_list = self._detect_word_list_item(document, paragraph)
        if word_list:
            ordered, indent = word_list
            return DocumentBlock(
                type="list_item",
                text=text,
                ordered=ordered,
                indent=indent,
            )

        # 3. Lista por estilo (fallback real)
        if style_name.lower().strip() == "list paragraph":
            return DocumentBlock(
                type="list_item",
                text=text,
                ordered=False,
                indent=0,
            )

        # 4. Lista manual (•, 1), etc.)
        list_item = self._detect_list_item(paragraph.text)
        if list_item:
            _, ordered, clean_text, indent = list_item
            return DocumentBlock(
                type="list_item",
                text=clean_text,
                ordered=ordered,
                indent=indent,
            )

        # 5. Heading por regex
        heading_level = self._detect_heading_level_from_text(text)
        if heading_level is not None:
            return DocumentBlock(
                type="heading",
                text=text,
                level=heading_level,
            )

        # 6. Código
        if self._looks_like_code(paragraph):
            return DocumentBlock(
                type="code",
                text=paragraph.text,
                language=None,
            )

        # 7. Párrafo normal
        return DocumentBlock(
            type="paragraph",
            text=text,
        )

    def _looks_like_ordered_list_text(self, text: str) -> bool:
        normalized = text.lower().strip()

        ordered_starters = (
            "primer ",
            "segundo ",
            "tercer ",
            "cuarto ",
            "quinto ",
            "sexto ",
            "séptimo ",
            "septimo ",
            "octavo ",
            "noveno ",
            "décimo ",
            "decimo ",
        )

        return normalized.startswith(ordered_starters)

        list_item = self._detect_list_item(paragraph.text)
        if list_item:
            _, ordered, clean_text, indent = list_item
            return DocumentBlock(
                type="list_item",
                text=clean_text,
                ordered=ordered,
                indent=indent,
            )

        heading_level = self._detect_heading_level_from_text(text)
        if heading_level is not None:
            return DocumentBlock(
                type="heading",
                text=text,
                level=heading_level,
            )

        if self._looks_like_code(paragraph):
            return DocumentBlock(
                type="code",
                text=paragraph.text,
                language=None,
            )

        return DocumentBlock(
            type="paragraph",
            text=text,
        )

    def _table_from_element(
        self,
        document: Document,
        element: Any,
    ) -> DocumentBlock | None:
        table = self._lookup_table(document, element)

        if table is None:
            return None

        rows: list[list[str]] = []

        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return None

        return DocumentBlock(type="table", rows=rows)

    def _lookup_paragraph(self, document: Document, element: Any) -> Any | None:
        for paragraph in document.paragraphs:
            if paragraph._element is element:
                return paragraph
        return None

    def _lookup_table(self, document: Document, element: Any) -> Any | None:
        for table in document.tables:
            if table._element is element:
                return table
        return None

    def _detect_heading_level(self, style_name: str, text: str) -> int | None:
        style_level = self._detect_heading_level_from_style(style_name)
        if style_level is not None:
            return style_level

        return self._detect_heading_level_from_text(text)

    def _detect_heading_level_from_style(self, style_name: str) -> int | None:
        normalized = style_name.lower().strip()

        if normalized.startswith("heading"):
            parts = normalized.split()
            if len(parts) >= 2 and parts[-1].isdigit():
                return int(parts[-1])
            return 1

        if normalized.startswith("título") or normalized.startswith("titulo"):
            return 1

        return None

    def _detect_heading_level_from_text(self, text: str) -> int | None:
        normalized = text.strip()

        if not self._can_be_heading(normalized):
            return None

        if re.match(r"^\s*\d+[\.\)\-]\s+\S+", normalized):
            if len(normalized.split()) <= 6:
                return None

        for pattern, level in self.HEADING_PATTERNS:
            if pattern.match(normalized):
                return level

        if self._looks_like_short_title(normalized):
            return 2

        return None

    def _can_be_heading(self, text: str) -> bool:
        if not text:
            return False

        if len(text) > self.MAX_HEADING_LENGTH:
            return False

        if "\n" in text:
            return False

        if text.count(".") >= 2:
            return False

        words = text.split()
        if len(words) > 14:
            return False

        return True

    def _looks_like_short_title(self, text: str) -> bool:
        lower = text.lower().strip()

        common_titles = {
            "introducción",
            "introduccion",
            "conclusión",
            "conclusion",
            "conclusiones",
            "resumen",
            "abstract",
            "referencias",
            "bibliografía",
            "bibliografia",
            "metodología",
            "metodologia",
            "resultados",
            "discusión",
            "discusion",
            "marco teórico",
            "marco teorico",
            "estado del arte",
            "objetivos",
            "limitaciones",
            "trabajos futuros",
            "líneas futuras",
            "lineas futuras",
        }

        if lower in common_titles:
            return True

        words = text.split()

        if 2 <= len(words) <= 6:
            has_sentence_punctuation = text.endswith(".") or text.endswith(";")
            starts_upper = text[0].isupper()

            if starts_upper and not has_sentence_punctuation:
                return True

        return False

    def _looks_like_code(self, paragraph: Any) -> bool:
        style_name = paragraph.style.name.lower().strip() if paragraph.style else ""
        text = paragraph.text
        stripped = text.strip()

        if "code" in style_name or "source" in style_name:
            return True

        if any(
            marker in text
            for marker in (
                "def ",
                "class ",
                "return ",
                "import ",
                "from ",
                "{",
                "}",
                "</",
            )
        ):
            return True

        code_markers = (
            "def ",
            "class ",
            "import ",
            "from ",
            "public ",
            "private ",
            "function ",
            "const ",
            "let ",
            "var ",
            "{",
            "}",
            "</",
        )

        return stripped.startswith(code_markers)

    def _detect_list_item(self, text: str) -> tuple[bool, bool, str, int] | None:
        original = text.rstrip()
        stripped = original.strip()

        leading_spaces = len(original) - len(original.lstrip(" "))
        indent = max(leading_spaces // 2, 0)

        if stripped.startswith(("• ", "- ", "* ")):
            return (True, False, stripped[2:].strip(), indent)

        match = re.match(r"^\s*(\d+)[\.\)\-]\s+(.*)$", stripped)
        if match:
            return (True, True, match.group(2).strip(), indent)

        match = re.match(r"^\s*[a-zA-Z][\)\.]\s+(.*)$", stripped)
        if match:
            return (True, True, match.group(1).strip(), indent)

        return None

    def _detect_word_list_item(
        self,
        document: Document,
        paragraph: Any,
    ) -> tuple[bool, int] | None:
        paragraph_properties = paragraph._element.pPr

        if paragraph_properties is None or paragraph_properties.numPr is None:
            return None

        num_pr = paragraph_properties.numPr
        num_id_element = num_pr.numId
        ilvl_element = num_pr.ilvl

        if num_id_element is None:
            return None

        num_id = num_id_element.val
        indent = int(ilvl_element.val) if ilvl_element is not None else 0

        ordered = self._is_ordered_word_list(document, str(num_id), indent)

        return ordered, indent

    def _is_ordered_word_list(
        self,
        document: Document,
        num_id: str,
        level: int,
    ) -> bool:
        numbering = document.part.numbering_part.element

        abstract_num_id = None

        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) != num_id:
                continue

            abstract = num.find(qn("w:abstractNumId"))
            if abstract is not None:
                abstract_num_id = abstract.get(qn("w:val"))
                break

        if abstract_num_id is None:
            return False

        for abstract_num in numbering.findall(qn("w:abstractNum")):
            if abstract_num.get(qn("w:abstractNumId")) != abstract_num_id:
                continue

            for lvl in abstract_num.findall(qn("w:lvl")):
                if lvl.get(qn("w:ilvl")) != str(level):
                    continue

                num_fmt = lvl.find(qn("w:numFmt"))
                if num_fmt is None:
                    return False

                value = num_fmt.get(qn("w:val"))

                return value in {
                    "decimal",
                    "upperRoman",
                    "lowerRoman",
                    "upperLetter",
                    "lowerLetter",
                }

        return False