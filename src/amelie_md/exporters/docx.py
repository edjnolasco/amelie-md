from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from amelie_md.parsing.inline_parser import parse_inline

from amelie_md.styles.docx import AcademicDocxStyle


@dataclass(frozen=True)
class DocxMetadata:
    title: str | None = None
    author: str | None = None
    date: str | None = None


class DocxExporter:
    """Professional DOCX exporter for Amelie MD.

    v1.2:
    - Renders AmelieDocument directly.
    - Removes MarkdownIt as intermediate renderer.
    - Keeps academic styling, cover, pagination, TOC, tables and code blocks.
    """

    MANUAL_HEADING_NUMBER_RE = re.compile(
        r"^\s*(?:\d+(?:\.\d+)*\.?|[IVXLCDM]+\.|[A-Z]\.)\s+",
        re.IGNORECASE,
    )

    CODE_FONT_FAMILY = "Consolas"

    def __init__(
        self,
        metadata: DocxMetadata | None = None,
        style: str = "academic",
        style_spec: Any | None = None,
    ) -> None:
        self.metadata = metadata or DocxMetadata()
        self.style = style.lower().strip()
        self.style_spec = style_spec

        self.heading_num_id: int | None = None
        self._ordered_list_stack: list[int] = []

    def export(self, document: Any, output_path: str | Path) -> Path:
        """Backward-compatible public entry point.

        In v1.2, this method expects an AmelieDocument-like object with
        a `.blocks` attribute. Markdown strings are intentionally rejected
        because Markdown is no longer used as an intermediate representation.
        """
        if isinstance(document, str):
            raise TypeError(
                "DocxExporter.export() no longer accepts Markdown text in v1.2. "
                "Pass an AmelieDocument instance or use export_document(document, output_path)."
            )

        return self.export_document(document, output_path)

    def export_document(self, amelie_document: Any, output_path: str | Path) -> Path:
        """Render an AmelieDocument directly to DOCX."""
        path = Path(output_path)
        document = Document()

        self._configure_document(document)
        self._apply_style(document)
        self._apply_style_spec(document, self.style_spec)
        self._apply_academic_spacing(document)

        self.heading_num_id = self._ensure_heading_numbering(document)
        self._ordered_list_stack = []

        if self._has_cover():
            self._add_cover(document)
            self._configure_content_section_pagination(document)

        blocks = self._document_blocks(amelie_document)

        for block in blocks:
            self._render_block(document, block)

        if self.style_spec:
            self._apply_style_spec(document, self.style_spec)
            self._apply_academic_spacing(document)
            self._apply_style_spec_to_existing_content(document)

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
        return path

    def _document_blocks(self, amelie_document: Any) -> list[Any]:
        blocks = self._value(amelie_document, "blocks", default=None)

        if blocks is None:
            raise TypeError(
                "export_document() expects an AmelieDocument-like object "
                "with a `.blocks` attribute."
            )

        return list(blocks)

    def _render_block(self, document: Document, block: Any) -> None:
        block_type = self._block_type(block)

        # Defensive cleanup: avoid rendering empty structural residues,
        # especially empty list items that appear as bullets in TOC/PDF output.
        if block_type in {"heading", "paragraph", "list_item"}:
            text = self._block_text(block).strip()
            if not text:
                return

        if block_type == "code":
            code = str(
                self._value(block, "code", default=None)
                or self._value(block, "content", default=None)
                or self._value(block, "text", default="")
            ).strip()
            if not code:
                return

        if block_type == "table":
            rows = self._table_rows(block)
            if not rows:
                return

        if block_type != "list_item":
            self._ordered_list_stack = []

        if block_type == "heading":
            self._render_heading_block(document, block)
            return

        if block_type == "paragraph":
            self._render_paragraph_block(document, block)
            return

        if block_type == "list_item":
            self._render_list_item_block(document, block)
            return

        if block_type == "table":
            self._render_table_block(document, block)
            return

        if block_type == "code":
            self._render_code_block(document, block)
            return

        if block_type == "toc":
            self._add_toc(document)
            return

        # Safe fallback: unknown blocks become paragraphs if they expose text.
        text = self._block_text(block).strip()
        if text:
            paragraph = document.add_paragraph(style="Normal")
            self._add_plain_text_runs(paragraph, text)

    def _block_type(self, block: Any) -> str:
        value = (
            self._value(block, "type", default=None)
            or self._value(block, "kind", default=None)
            or self._value(block, "block_type", default=None)
        )

        return str(value or "").lower().strip()

    def _render_heading_block(self, document: Document, block: Any) -> None:
        level = int(self._value(block, "level", default=1) or 1)
        level = max(1, min(level, 3))

        text = self._strip_manual_heading_number(self._block_text(block))

        paragraph = document.add_heading(level=level)

        if self.heading_num_id is not None:
            self._apply_heading_numbering(paragraph, level, self.heading_num_id)

        self._add_plain_text_runs(paragraph, text)

    def _render_paragraph_block(self, document: Document, block: Any) -> None:
        text = self._block_text(block)

        if text.strip() == "[[TOC]]":
            self._add_toc(document)
            return

        paragraph = document.add_paragraph(style="Normal")
        self._add_plain_text_runs(paragraph, text)

        if not paragraph.text.strip():
            self._remove_empty_paragraph(paragraph)

    def _render_list_item_block(self, document: Document, block: Any) -> None:
        text = self._block_text(block)
        ordered = bool(self._value(block, "ordered", default=False))
        level = int(
            self._value(block, "level", default=None)
            if self._value(block, "level", default=None) is not None
            else self._value(block, "indent", default=0)
        )
        level = max(level, 0)

        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.32 * (level + 1))
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)

        if ordered:
            marker = self._ordered_marker(self._next_ordered_numbers(level))
        else:
            marker = self._bullet_marker(level)

        marker_run = paragraph.add_run(marker)
        marker_run.bold = False
        self._apply_inline_style_spec(marker_run)

        self._add_plain_text_runs(paragraph, text)

    def _next_ordered_numbers(self, level: int) -> tuple[int, ...]:
        while len(self._ordered_list_stack) <= level:
            self._ordered_list_stack.append(0)

        self._ordered_list_stack[level] += 1
        del self._ordered_list_stack[level + 1 :]

        return tuple(self._ordered_list_stack[: level + 1])

    def _ordered_marker(self, numbers: tuple[int, ...]) -> str:
        if not numbers:
            return "1. "

        return ".".join(str(number) for number in numbers) + ". "

    def _bullet_marker(self, level: int) -> str:
        markers = ("• ", "◦ ", "▪ ")
        return markers[level % len(markers)]

    def _render_table_block(self, document: Document, block: Any) -> None:
        rows = self._table_rows(block)

        if not rows:
            return

        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        table.autofit = True

        for row_index, row in enumerate(rows):
            for column_index in range(column_count):
                value = row[column_index] if column_index < len(row) else ""
                cell = table.cell(row_index, column_index)
                cell.text = str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                self._set_cell_margins(cell)

                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif self._looks_numeric(str(value)):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in paragraph.runs:
                        self._apply_run_font_from_spec(run)

                        if row_index == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

                if row_index == 0:
                    self._set_cell_shading(cell, "1F4E79")

        document.add_paragraph()

    def _table_rows(self, block: Any) -> list[list[str]]:
        rows = self._value(block, "rows", default=None)

        if rows is None:
            rows = self._value(block, "data", default=None)

        if rows is None:
            return []

        normalized_rows: list[list[str]] = []

        for row in rows:
            if isinstance(row, dict):
                cells = row.get("cells", [])
            else:
                cells = self._value(row, "cells", default=row)

            normalized_rows.append([str(self._cell_text(cell)) for cell in cells])

        return normalized_rows

    def _cell_text(self, cell: Any) -> str:
        if isinstance(cell, str):
            return cell

        if isinstance(cell, (int, float)):
            return str(cell)

        return str(
            self._value(cell, "text", default=None)
            or self._value(cell, "content", default="")
        )

    def _render_code_block(self, document: Document, block: Any) -> None:
        code = (
            self._value(block, "code", default=None)
            or self._value(block, "content", default=None)
            or self._value(block, "text", default="")
        )

        self._add_code_block(document, str(code))

    def _block_text(self, block: Any) -> str:
        value = (
            self._value(block, "text", default=None)
            or self._value(block, "content", default=None)
            or self._value(block, "value", default="")
        )

        return str(value or "")

    def _value(self, obj: Any, name: str, default: Any = None) -> Any:
        if obj is None:
            return default

        if isinstance(obj, dict):
            return obj.get(name, default)

        return getattr(obj, name, default)

    def _add_plain_text_runs(self, paragraph: Any, text: str) -> None:
        lines = str(text).splitlines()

        if not lines:
            return

        for index, line in enumerate(lines):
            if index > 0:
                paragraph.add_run().add_break()

            if not line:
                continue

            for inline_run in parse_inline(line):
                run = paragraph.add_run(inline_run.text)
                self._apply_inline_style_spec(run)

                if inline_run.bold:
                    run.bold = True

                if inline_run.italic:
                    run.italic = True

                if inline_run.code:
                    self._apply_code_run_style(run, block=False)

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        if self.metadata.title:
            document.core_properties.title = self.metadata.title
        if self.metadata.author:
            document.core_properties.author = self.metadata.author

    def _apply_style(self, document: Document) -> None:
        if self.style != "academic":
            raise ValueError(f"Unsupported DOCX style: {self.style}")
        AcademicDocxStyle().apply(document)

    def _style_value(self, style_spec: Any | None, name: str, default: Any = None) -> Any:
        if style_spec is None:
            return default

        if isinstance(style_spec, dict):
            return style_spec.get(name, default)

        if hasattr(style_spec, name):
            return getattr(style_spec, name)

        if hasattr(style_spec, "model_dump"):
            return style_spec.model_dump().get(name, default)

        if hasattr(style_spec, "dict"):
            return style_spec.dict().get(name, default)

        return default

    def _style_has_value(self, style_spec: Any | None, name: str) -> bool:
        return self._style_value(style_spec, name) not in {None, ""}

    def _apply_style_spec(self, document: Document, style_spec: Any | None) -> None:
        if not style_spec:
            return

        self._apply_font_spec(document, style_spec)
        self._apply_spacing_spec(document, style_spec)
        self._apply_page_layout_spec(document, style_spec)

    def _apply_font_spec(self, document: Document, style_spec: Any) -> None:
        def set_font(style: Any) -> None:
            font = style.font

            font_family = self._style_value(style_spec, "font_family")
            font_size = self._style_value(style_spec, "font_size")

            if font_family:
                self._set_font_family(font, font_family)

            if font_size:
                font.size = Pt(font_size)

        style_names = ["Normal", "Body Text", "Caption", "Table Grid"]

        for i in range(1, 10):
            style_names.append(f"Heading {i}")

        for style_name in style_names:
            if style_name in document.styles:
                set_font(document.styles[style_name])

    def _apply_spacing_spec(self, document: Document, style_spec: Any) -> None:
        spacing = self._style_value(style_spec, "spacing")

        if not spacing:
            return

        spacing_rule = None

        if spacing == "single":
            spacing_rule = WD_LINE_SPACING.SINGLE
        elif spacing == "double":
            spacing_rule = WD_LINE_SPACING.DOUBLE

        if spacing_rule is None:
            return

        style_names = ["Normal", "Body Text"]

        for i in range(1, 10):
            style_names.append(f"Heading {i}")

        for style_name in style_names:
            if style_name in document.styles:
                paragraph_format = document.styles[style_name].paragraph_format
                paragraph_format.line_spacing_rule = spacing_rule

    def _apply_academic_spacing(self, document: Document) -> None:
        has_explicit_spacing = bool(
            self.style_spec and self._style_value(self.style_spec, "spacing")
        )

        if "Normal" in document.styles:
            normal = document.styles["Normal"].paragraph_format
            normal.space_before = Pt(0)
            normal.space_after = Pt(6)
            normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if not has_explicit_spacing:
                normal.line_spacing = 1.15

        if "Body Text" in document.styles:
            body = document.styles["Body Text"].paragraph_format
            body.space_before = Pt(0)
            body.space_after = Pt(6)
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if not has_explicit_spacing:
                body.line_spacing = 1.15

        heading_spacing = {
            1: (18, 6),
            2: (14, 6),
            3: (12, 6),
            4: (10, 4),
            5: (8, 4),
            6: (8, 4),
            7: (8, 4),
            8: (8, 4),
            9: (8, 4),
        }

        for level, (before, after) in heading_spacing.items():
            style_name = f"Heading {level}"

            if style_name not in document.styles:
                continue

            paragraph_format = document.styles[style_name].paragraph_format
            paragraph_format.space_before = Pt(before)
            paragraph_format.space_after = Pt(after)
            paragraph_format.keep_with_next = True

            if not has_explicit_spacing:
                paragraph_format.line_spacing = 1.15

        if "Amelie Code Block" in document.styles:
            code_format = document.styles["Amelie Code Block"].paragraph_format
            code_format.space_before = Pt(0)
            code_format.space_after = Pt(0)
            code_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            code_font = document.styles["Amelie Code Block"].font
            self._set_font_family(code_font, self.CODE_FONT_FAMILY)
            code_font.size = Pt(9)

        if "Table Grid" in document.styles:
            table_font = document.styles["Table Grid"].font

            font_family = self._style_value(self.style_spec, "font_family")
            font_size = self._style_value(self.style_spec, "font_size")

            if font_family:
                self._set_font_family(table_font, font_family)

            if font_size:
                table_font.size = Pt(max(font_size - 1, 9))
            else:
                table_font.size = Pt(10)

    def _apply_page_layout_spec(self, document: Document, style_spec: Any) -> None:
        for section in document.sections:
            self._apply_section_layout(section, style_spec)

    def _to_inches(self, value: str | None, default: float) -> Any:
        if not value:
            return Inches(default)

        normalized = str(value).lower().strip().replace('"', "").replace("inches", "in")
        normalized = (
            normalized.replace("inch", "in")
            .replace("pulgadas", "in")
            .replace("pulgada", "in")
        )

        if normalized.endswith("in"):
            normalized = normalized.removesuffix("in").strip()

        try:
            return Inches(float(normalized))
        except ValueError:
            return Inches(default)

    def _has_cover(self) -> bool:
        return bool(self.metadata.title or self.metadata.author or self.metadata.date)

    def _add_cover(self, document: Document) -> None:
        title = self.metadata.title or "Untitled Document"
        author = self.metadata.author or ""
        document_date = self.metadata.date or date.today().isoformat()

        document.add_paragraph("\n" * 5)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(28)

        font_family = self._style_value(self.style_spec, "font_family")
        if font_family:
            self._set_font_family(run.font, font_family)

        document.add_paragraph("\n" * 2)

        if author:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(author)
            run.font.size = Pt(14)

            if font_family:
                self._set_font_family(run.font, font_family)

        document.add_paragraph("\n")

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(document_date)
        run.font.size = Pt(12)

        if font_family:
            self._set_font_family(run.font, font_family)

        document.add_section(WD_SECTION_START.NEW_PAGE)

    def _configure_content_section_pagination(self, document: Document) -> None:
        if len(document.sections) < 2:
            return

        cover_section = document.sections[0]
        content_section = document.sections[1]

        cover_section.footer.is_linked_to_previous = False
        self._clear_footer(cover_section)

        content_section.footer.is_linked_to_previous = False
        self._restart_page_numbering(content_section, start=1)
        self._add_page_number(content_section)

        if self.style_spec:
            self._apply_section_layout(content_section, self.style_spec)

    def _apply_section_layout(self, section: Any, style_spec: Any) -> None:
        page_size = self._style_value(style_spec, "page_size")

        if page_size == "letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        section.top_margin = self._to_inches(
            self._style_value(style_spec, "margin_top"),
            default=1,
        )
        section.bottom_margin = self._to_inches(
            self._style_value(style_spec, "margin_bottom"),
            default=1,
        )
        section.left_margin = self._to_inches(
            self._style_value(style_spec, "margin_left"),
            default=1,
        )
        section.right_margin = self._to_inches(
            self._style_value(style_spec, "margin_right"),
            default=1,
        )

    def _restart_page_numbering(self, section: Any, start: int = 1) -> None:
        sect_pr = section._sectPr

        existing_pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if existing_pg_num_type is not None:
            sect_pr.remove(existing_pg_num_type)

        page_number_type = OxmlElement("w:pgNumType")
        page_number_type.set(qn("w:start"), str(start))
        sect_pr.append(page_number_type)

    def _add_page_number(self, section: Any) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._clear_paragraph(paragraph)

        paragraph.add_run("Página ")
        self._add_word_field(paragraph, "PAGE")
        paragraph.add_run(" de ")
        self._add_word_field(paragraph, "SECTIONPAGES")

    def _add_word_field(self, paragraph: Any, instruction_text: str) -> None:
        run = paragraph.add_run()

        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")

        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = instruction_text

        field_end = OxmlElement("w:fldChar")
        field_end.set(qn("w:fldCharType"), "end")

        run._r.append(field_begin)
        run._r.append(instruction)
        run._r.append(field_end)

    def _clear_footer(self, section: Any) -> None:
        footer = section.footer

        for paragraph in footer.paragraphs:
            self._clear_paragraph(paragraph)

    def _clear_paragraph(self, paragraph: Any) -> None:
        for run in list(paragraph.runs):
            paragraph._element.remove(run._element)

    def _apply_run_font_from_spec(self, run: Any) -> None:
        font_family = self._style_value(self.style_spec, "font_family")
        font_size = self._style_value(self.style_spec, "font_size")

        if font_family:
            self._set_font_family(run.font, font_family)
        elif not self.style_spec:
            self._set_font_family(run.font, "Calibri")

        if font_size:
            run.font.size = Pt(font_size)
        elif not self.style_spec:
            run.font.size = Pt(10)

    def _ensure_heading_numbering(self, document: Document) -> int:
        numbering = document.part.numbering_part.element

        abstract_id = str(self._next_numbering_id(numbering, "abstractNum"))
        num_id = str(self._next_numbering_id(numbering, "num"))

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abstract_id)

        multi_level_type = OxmlElement("w:multiLevelType")
        multi_level_type.set(qn("w:val"), "multilevel")
        abstract.append(multi_level_type)

        for ilvl, level_text in {
            0: "%1",
            1: "%1.%2",
            2: "%1.%2.%3",
        }.items():
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(ilvl))

            level.append(self._xml("w:start", val="1"))
            level.append(self._xml("w:numFmt", val="decimal"))
            level.append(self._xml("w:lvlText", val=level_text))
            level.append(self._xml("w:lvlJc", val="left"))

            paragraph_properties = OxmlElement("w:pPr")
            indent = OxmlElement("w:ind")
            indent.set(qn("w:left"), str(360 * ilvl))
            indent.set(qn("w:hanging"), "0")
            paragraph_properties.append(indent)
            level.append(paragraph_properties)

            abstract.append(level)

        numbering.append(abstract)

        numbering_instance = OxmlElement("w:num")
        numbering_instance.set(qn("w:numId"), num_id)
        numbering_instance.append(self._xml("w:abstractNumId", val=abstract_id))
        numbering.append(numbering_instance)

        return int(num_id)

    def _next_numbering_id(self, root: Any, element_name: str) -> int:
        existing_ids: list[int] = []

        for element in root.findall(qn(f"w:{element_name}")):
            attribute = "w:abstractNumId" if element_name == "abstractNum" else "w:numId"
            value = element.get(qn(attribute))

            if value is not None and value.isdigit():
                existing_ids.append(int(value))

        return max(existing_ids, default=0) + 1

    def _apply_heading_numbering(self, paragraph: Any, level: int, num_id: int) -> None:
        paragraph_properties = paragraph._element.get_or_add_pPr()

        existing_numbering = paragraph_properties.find(qn("w:numPr"))
        if existing_numbering is not None:
            paragraph_properties.remove(existing_numbering)

        numbering_properties = OxmlElement("w:numPr")
        numbering_properties.append(self._xml("w:ilvl", val=str(level - 1)))
        numbering_properties.append(self._xml("w:numId", val=str(num_id)))

        paragraph_properties.append(numbering_properties)

    def _add_toc(self, document: Document) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Table of Contents")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)

        font_family = self._style_value(self.style_spec, "font_family")
        if font_family:
            self._set_font_family(run.font, font_family)

        paragraph = document.add_paragraph()
        field_run = paragraph.add_run()

        field_run._r.append(self._xml("w:fldChar", fldCharType="begin"))

        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
        field_run._r.append(instruction)

        field_run._r.append(self._xml("w:fldChar", fldCharType="separate"))

        placeholder = paragraph.add_run(
            "Right-click and update field to generate table of contents"
        )
        placeholder.italic = True

        end_run = paragraph.add_run()
        end_run._r.append(self._xml("w:fldChar", fldCharType="end"))

        document.add_paragraph()
        document.add_page_break()

    def _xml(self, tag: str, **attrs: str) -> Any:
        element = OxmlElement(tag)

        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), value)

        return element

    def _set_cell_margins(
        self,
        cell: Any,
        top: int = 120,
        start: int = 120,
        bottom: int = 120,
        end: int = 120,
    ) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")

        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)

        for margin_name, margin_value in {
            "top": top,
            "start": start,
            "bottom": bottom,
            "end": end,
        }.items():
            node = tc_mar.find(qn(f"w:{margin_name}"))

            if node is None:
                node = OxmlElement(f"w:{margin_name}")
                tc_mar.append(node)

            node.set(qn("w:w"), str(margin_value))
            node.set(qn("w:type"), "dxa")

    def _set_cell_shading(self, cell: Any, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))

        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)

        shading.set(qn("w:fill"), fill)

    def _looks_numeric(self, value: str) -> bool:
        cleaned = value.strip().replace(",", "").replace("%", "")

        if not cleaned:
            return False

        try:
            float(cleaned)
        except ValueError:
            return False

        return True

    def _add_code_block(self, document: Document, code: str) -> None:
        lines = code.rstrip("\n").splitlines() or [""]

        for line in lines:
            paragraph = document.add_paragraph(style="Amelie Code Block")
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            run = paragraph.add_run(line)
            self._apply_code_run_style(run, block=True)

    def _strip_manual_heading_number(self, text: str) -> str:
        return self.MANUAL_HEADING_NUMBER_RE.sub("", text, count=1)

    def _apply_inline_style_spec(self, run: Any) -> None:
        font_family = self._style_value(self.style_spec, "font_family")
        font_size = self._style_value(self.style_spec, "font_size")

        if font_family:
            self._set_font_family(run.font, font_family)

        if font_size:
            run.font.size = Pt(font_size)

    def _apply_code_run_style(self, run: Any, block: bool) -> None:
        self._set_font_family(run.font, self.CODE_FONT_FAMILY)

        font_size = self._style_value(self.style_spec, "font_size")

        if font_size:
            size = max(font_size - 1, 8) if block else max(font_size - 2, 8)
            run.font.size = Pt(size)
        else:
            run.font.size = Pt(9)

        run.font.color.rgb = RGBColor(45, 45, 45)

    def _apply_style_spec_to_existing_content(self, document: Document) -> None:
        font_family = self._style_value(self.style_spec, "font_family")
        font_size = self._style_value(self.style_spec, "font_size")

        if not font_family and not font_size:
            return

        for paragraph in document.paragraphs:
            is_code_paragraph = paragraph.style and paragraph.style.name == "Amelie Code Block"

            for run in paragraph.runs:
                if is_code_paragraph:
                    self._apply_code_run_style(run, block=True)
                    continue

                if font_family:
                    self._set_font_family(run.font, font_family)

                if font_size:
                    run.font.size = Pt(font_size)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if font_family:
                                self._set_font_family(run.font, font_family)

                            if font_size:
                                run.font.size = Pt(max(font_size - 1, 9))

    def _set_font_family(self, font: Any, family: str) -> None:
        font.name = family

        if font.element.rPr is None:
            font.element.get_or_add_rPr()

        font.element.rPr.rFonts.set(qn("w:ascii"), family)
        font.element.rPr.rFonts.set(qn("w:hAnsi"), family)
        font.element.rPr.rFonts.set(qn("w:eastAsia"), family)
        font.element.rPr.rFonts.set(qn("w:cs"), family)

    def _remove_empty_paragraph(self, paragraph: Any) -> None:
        element = paragraph._element
        element.getparent().remove(element)